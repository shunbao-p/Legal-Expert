from __future__ import annotations

import datetime as dt
import logging
import re
import threading
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional, TYPE_CHECKING

from langchain_core.documents import Document

if TYPE_CHECKING:
    from main import AdvancedGraphRAGSystem

logger = logging.getLogger(__name__)


@dataclass
class SessionFileRecord:
    file_id: str
    chat_id: str
    file_name: str
    modality: str
    size_bytes: int
    uploaded_at: str
    active: bool = True
    status: str = "ready"
    error: str = ""
    local_path: str = ""
    chunks: list[Document] = field(default_factory=list)


class RAGDemoService:
    """Thin service wrapper that exposes the existing GraphRAG system to FastAPI."""

    def __init__(self) -> None:
        self._lock = threading.Lock()
        self._system: Optional["AdvancedGraphRAGSystem"] = None
        self._initialized = False
        self._startup_error = ""
        self._chat_sessions: set[str] = set()
        self._chat_files: dict[str, dict[str, SessionFileRecord]] = {}
        self._upload_root = Path("data/session_uploads")
        self._upload_root.mkdir(parents=True, exist_ok=True)

    @property
    def startup_error(self) -> str:
        return self._startup_error

    @property
    def initialized(self) -> bool:
        return self._initialized

    @property
    def system_ready(self) -> bool:
        return bool(self._system and self._system.system_ready)

    def startup(self) -> None:
        with self._lock:
            if self._initialized and self._system is not None:
                return

            try:
                from main import AdvancedGraphRAGSystem

                system = AdvancedGraphRAGSystem()
                system.initialize_system()
                system.build_knowledge_base()
                self._system = system
                self._initialized = True
                self._startup_error = ""
                logger.info("RAG demo service startup complete")
            except Exception as exc:
                self._startup_error = str(exc)
                logger.exception("RAG demo service startup failed")
                raise

    def shutdown(self) -> None:
        with self._lock:
            if self._system is not None:
                self._system._cleanup()
            self._system = None
            self._initialized = False
            self._chat_sessions.clear()
            self._chat_files.clear()

    def health(self) -> dict:
        if self.system_ready:
            status = "ready"
        elif self._startup_error:
            status = "failed"
        else:
            status = "starting"
        return {
            "status": status,
            "initialized": self._initialized,
            "system_ready": self.system_ready,
            "startup_error": self._startup_error,
        }

    def create_chat_session(self) -> str:
        chat_id = uuid.uuid4().hex
        with self._lock:
            self._chat_sessions.add(chat_id)
            self._chat_files.setdefault(chat_id, {})
        return chat_id

    def _assert_chat_session_exists(self, chat_id: str) -> None:
        with self._lock:
            if chat_id not in self._chat_sessions:
                raise KeyError("会话不存在或已失效，请刷新页面后重试。")

    @staticmethod
    def _infer_modality(file_name: str, content_type: str) -> str:
        suffix = Path(file_name).suffix.lower()
        if suffix == ".pdf" or "pdf" in content_type:
            return "pdf"
        if suffix in {".docx", ".doc"}:
            return "word"
        if suffix in {".xlsx", ".xls", ".csv"}:
            return "excel"
        if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".webp"} or content_type.startswith("image/"):
            return "image"
        if suffix in {".mp3", ".wav", ".m4a", ".flac"} or content_type.startswith("audio/"):
            return "audio"
        return "text"

    @staticmethod
    def _safe_name(file_name: str) -> str:
        base = Path(file_name or "upload.bin").name
        safe = re.sub(r"[^0-9A-Za-z_\-.()\u4e00-\u9fff]+", "_", base)
        return safe[:120] or "upload.bin"

    @staticmethod
    def _chunk_text(text: str, chunk_size: int = 700, overlap: int = 120) -> list[str]:
        compact = str(text or "").strip()
        if not compact:
            return []
        chunks: list[str] = []
        start = 0
        step = max(1, chunk_size - overlap)
        while start < len(compact):
            end = min(len(compact), start + chunk_size)
            part = compact[start:end].strip()
            if part:
                chunks.append(part)
            if end >= len(compact):
                break
            start += step
        return chunks

    @staticmethod
    def _keyword_score(question: str, content: str) -> float:
        query = str(question or "").strip().lower()
        text = str(content or "").lower()
        if not query or not text:
            return 0.0
        terms = [term for term in re.split(r"\s+", query) if term]
        if len(terms) <= 1:
            cjk_terms = sorted({ch for ch in query if "\u4e00" <= ch <= "\u9fff"})
            if cjk_terms:
                terms = cjk_terms
        if not terms:
            return 0.0
        hits = sum(1 for term in terms if term in text)
        return round(hits / float(len(terms)), 4)

    @staticmethod
    def _is_file_content_question(question: str) -> bool:
        text = str(question or "").strip().lower()
        if not text:
            return False
        patterns = [
            "文件内容",
            "文档内容",
            "根据文件",
            "基于文件",
            "回答文件",
            "这个文件",
            "该文件",
            "文件讲了什么",
            "总结文件",
            "读取文件",
        ]
        return any(pattern in text for pattern in patterns)

    def _extract_text(self, local_path: Path, modality: str, raw_bytes: bytes) -> tuple[str, str]:
        if modality == "pdf":
            try:
                from pypdf import PdfReader

                reader = PdfReader(str(local_path))
                return "\n".join((page.extract_text() or "") for page in reader.pages).strip(), ""
            except ModuleNotFoundError:
                return "", "缺少依赖 pypdf，当前无法解析 PDF。"
            except Exception as exc:
                logger.warning("PDF解析失败，回退到二进制解码: %s", exc)
                return "", f"PDF 解析失败: {exc}"
        elif modality == "word":
            try:
                from docx import Document as DocxDocument

                document = DocxDocument(str(local_path))
                return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text).strip(), ""
            except ModuleNotFoundError:
                return "", "缺少依赖 python-docx，当前无法解析 Word。"
            except Exception as exc:
                logger.warning("Word解析失败: %s", exc)
                return "", f"Word 解析失败: {exc}"
        elif modality == "excel":
            try:
                if local_path.suffix.lower() == ".csv":
                    return raw_bytes.decode("utf-8", errors="ignore"), ""
                from openpyxl import load_workbook

                wb = load_workbook(str(local_path), read_only=True, data_only=True)
                lines: list[str] = []
                for sheet in wb.worksheets:
                    lines.append(f"[sheet]{sheet.title}")
                    for row in sheet.iter_rows(values_only=True):
                        values = [str(item).strip() for item in row if item is not None and str(item).strip()]
                        if values:
                            lines.append(" | ".join(values))
                return "\n".join(lines).strip(), ""
            except ModuleNotFoundError:
                return "", "缺少依赖 openpyxl，当前无法解析 Excel。"
            except Exception as exc:
                logger.warning("Excel解析失败: %s", exc)
                return "", f"Excel 解析失败: {exc}"
        elif modality == "image":
            try:
                from PIL import Image
                import pytesseract

                text = pytesseract.image_to_string(Image.open(str(local_path)), lang="chi_sim+eng")
                return text.strip(), ""
            except ModuleNotFoundError:
                return "", "缺少依赖 Pillow/pytesseract，当前无法执行图片 OCR。"
            except Exception as exc:
                logger.warning("图片OCR不可用或失败: %s", exc)
                return "", f"图片 OCR 失败: {exc}"
        elif modality == "audio":
            # MVP阶段占位：后续接ASR模块（faster-whisper/云ASR）
            return "", "音频转写尚未启用，待接入 ASR。"
        return raw_bytes.decode("utf-8", errors="ignore").strip(), ""

    def upload_file(self, chat_id: str, file_name: str, content_type: str, data: bytes) -> dict[str, Any]:
        self._assert_chat_session_exists(chat_id)
        safe_name = self._safe_name(file_name)
        modality = self._infer_modality(safe_name, content_type or "")
        file_id = uuid.uuid4().hex
        uploaded_at = dt.datetime.now(dt.timezone.utc).isoformat()
        local_path = self._upload_root / f"{chat_id}_{file_id}_{safe_name}"
        local_path.parent.mkdir(parents=True, exist_ok=True)
        local_path.write_bytes(data)

        text, parse_error = self._extract_text(local_path, modality, data)
        chunks = self._chunk_text(text)
        documents: list[Document] = []
        for index, chunk_text in enumerate(chunks):
            chunk_id = f"{file_id}_chunk_{index}"
            documents.append(
                Document(
                    page_content=chunk_text,
                    metadata={
                        "chunk_id": chunk_id,
                        "display_title": safe_name,
                        "law_name": "",
                        "article_id": "",
                        "article_title": "",
                        "search_type": "session_file",
                        "search_source": "session_file",
                        "route_strategy": "session_file",
                        "chat_id": chat_id,
                        "file_id": file_id,
                        "modality": modality,
                    },
                )
            )

        status = "ready" if documents else ("error" if parse_error else "empty")
        error = "" if documents else (parse_error or "文件解析后无可用文本，可更换文件或补充OCR/ASR依赖。")
        record = SessionFileRecord(
            file_id=file_id,
            chat_id=chat_id,
            file_name=safe_name,
            modality=modality,
            size_bytes=len(data or b""),
            uploaded_at=uploaded_at,
            status=status,
            error=error,
            local_path=str(local_path),
            chunks=documents,
        )
        with self._lock:
            self._chat_files.setdefault(chat_id, {})[file_id] = record
        return {"file": self._file_to_dict(record)}

    def _file_to_dict(self, record: SessionFileRecord) -> dict[str, Any]:
        return {
            "file_id": record.file_id,
            "file_name": record.file_name,
            "modality": record.modality,
            "status": record.status,
            "size_bytes": record.size_bytes,
            "uploaded_at": record.uploaded_at,
            "active": record.active,
            "parsed_chunks": len(record.chunks),
            "error": record.error,
        }

    def list_chat_files(self, chat_id: str) -> dict[str, Any]:
        self._assert_chat_session_exists(chat_id)
        with self._lock:
            records = list(self._chat_files.get(chat_id, {}).values())
        records.sort(key=lambda item: item.uploaded_at, reverse=True)
        return {
            "chat_id": chat_id,
            "files": [self._file_to_dict(item) for item in records],
        }

    def delete_chat_file(self, chat_id: str, file_id: str) -> dict[str, Any]:
        self._assert_chat_session_exists(chat_id)
        record: Optional[SessionFileRecord] = None
        with self._lock:
            record = self._chat_files.get(chat_id, {}).pop(file_id, None)
        if record is None:
            raise KeyError("文件不存在或已删除。")
        try:
            if record.local_path:
                Path(record.local_path).unlink(missing_ok=True)
        except Exception:
            logger.exception("删除临时文件失败: %s", record.local_path)
        return {"chat_id": chat_id, "file_id": file_id, "deleted": True}

    def _resolve_active_file_ids(self, chat_id: str, active_file_ids: Optional[list[str]]) -> list[str]:
        with self._lock:
            all_records = self._chat_files.get(chat_id, {})
            if active_file_ids is None:
                return [fid for fid, rec in all_records.items() if rec.active and rec.status == "ready"]
            selected = []
            for file_id in active_file_ids:
                rec = all_records.get(file_id)
                if rec and rec.active and rec.status == "ready":
                    selected.append(file_id)
            return selected

    def _retrieve_session_documents(
        self,
        chat_id: str,
        question: str,
        active_file_ids: Optional[list[str]],
        top_k: int = 4,
    ) -> list[Document]:
        resolved_file_ids = self._resolve_active_file_ids(chat_id, active_file_ids)
        if not resolved_file_ids:
            return []

        candidates: list[tuple[float, Document]] = []
        fallback_chunks: list[Document] = []
        with self._lock:
            file_map = self._chat_files.get(chat_id, {})
            for file_id in resolved_file_ids:
                record = file_map.get(file_id)
                if not record:
                    continue
                if record.chunks:
                    fallback_chunks.append(record.chunks[0])
                for chunk_doc in record.chunks:
                    score = self._keyword_score(question, chunk_doc.page_content)
                    if score <= 0:
                        continue
                    chunk_doc.metadata["relevance_score"] = score
                    chunk_doc.metadata["final_score"] = score
                    candidates.append((score, chunk_doc))
        if not candidates:
            if not self._is_file_content_question(question):
                return []
            selected = fallback_chunks[:top_k]
            for doc in selected:
                doc.metadata["relevance_score"] = 0.05
                doc.metadata["final_score"] = 0.05
                doc.metadata["search_type"] = "session_file_fallback"
            return selected
        candidates.sort(key=lambda item: item[0], reverse=True)
        return [item[1] for item in candidates[:top_k]]

    def chat(
        self,
        chat_id: str,
        question: str,
        explain_routing: bool = False,
        active_file_ids: Optional[list[str]] = None,
    ) -> dict:
        self._assert_chat_session_exists(chat_id)
        if not self._initialized or self._system is None:
            self.startup()
        assert self._system is not None
        session_docs = self._retrieve_session_documents(chat_id, question, active_file_ids=active_file_ids)
        return self._system.ask_question_payload(
            question,
            explain_routing=explain_routing,
            chat_id=chat_id,
            active_file_ids=active_file_ids,
            prefetched_documents=session_docs,
        )
